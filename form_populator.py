"""
form_populator.py
InSync Tech — ARIA Intake Form Population

Takes the structured JSON from Claude Haiku extractor and generates
a fully populated version of the ARIA Client Intake Form as a .docx.
Checkboxes are checked (■) where the client selected them,
write-in fields are filled with the client's actual words.

Delivered alongside the summary PDF in the Resend email.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── CONSTANTS ───────────────────────────────────────────────────────────────
CHECKED   = "■"   # filled box — selected
UNCHECKED = "□"   # empty box  — not selected
BLUE      = RGBColor(0x1A, 0x3D, 0x6B)
LIGHT_BLUE= RGBColor(0x2E, 0x75, 0xB6)
GRAY      = RGBColor(0x40, 0x40, 0x40)
MID_GRAY  = RGBColor(0x66, 0x66, 0x66)
GREEN     = RGBColor(0x1E, 0x7E, 0x34)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)


# ─── HELPERS ─────────────────────────────────────────────────────────────────

def _safe(val, fallback="—"):
    if val is None:
        return fallback
    if isinstance(val, list):
        return ", ".join(str(v) for v in val) if val else fallback
    return str(val).strip() or fallback


def _matches(value, option: str) -> bool:
    """
    Fuzzy match — check if an option string appears in the extracted value.
    Works for lists of strings or a single string.
    """
    if value is None:
        return False
    option_lower = option.lower()
    if isinstance(value, list):
        return any(option_lower in str(v).lower() for v in value)
    return option_lower in str(value).lower()


def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_run(para, text, bold=False, italic=False, size=10,
             color=None, font="Arial"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def _section_header(doc, num, title):
    """Section header with bottom border simulation via paragraph border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    _add_run(p, f"{num}  ", bold=True, size=13, color=LIGHT_BLUE)
    _add_run(p, title.upper(), bold=True, size=13, color=BLUE)
    # Bottom border via paragraph XML
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _question_label(doc, text, required=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    _add_run(p, text, bold=True, size=10, color=GRAY)
    if required:
        _add_run(p, "  *required", italic=True, size=8, color=RGBColor(0xCC, 0x00, 0x00))


def _write_in(doc, value, hint="", lines=1):
    """Filled write-in box — shows the client's answer."""
    filled = _safe(value)
    # Build a 1-cell table as the box
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "FAFBFC")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if filled and filled != "—":
        _add_run(p, filled, size=10, color=GRAY)
    else:
        _add_run(p, hint if hint else "Not provided", italic=True, size=9,
                 color=RGBColor(0xAA, 0xAA, 0xAA))
    # Add extra blank lines if multi-line
    for _ in range(lines - 1):
        cell.add_paragraph()
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def _check_banner(doc, include_check_all=True, all_checked=False):
    """CHECK ALL THAT APPLY banner row."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "EAF2FB")
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    _add_run(p, "  CHECK ALL THAT APPLY", bold=True, size=8, color=LIGHT_BLUE)
    if include_check_all:
        mark = CHECKED if all_checked else UNCHECKED
        _add_run(p, f"          {mark}  Check All", bold=True, size=8, color=LIGHT_BLUE)


def _check_grid(doc, items, value):
    """
    Two-column checkbox grid. Items is list of dicts with 'label' and optional 'sub'.
    Checks boxes that match extracted value.
    """
    filtered = [i for i in items if i.get("label")]
    # Pad to even number
    if len(filtered) % 2 != 0:
        filtered.append({"label": ""})

    pairs = [(filtered[i], filtered[i+1]) for i in range(0, len(filtered), 2)]

    table = doc.add_table(rows=len(pairs), cols=2)
    for row_idx, (left, right) in enumerate(pairs):
        for col_idx, item in enumerate([(left, 0), (right, 1)]):
            opt, col = item
            cell = table.cell(row_idx, col)
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

            if not opt.get("label"):
                _add_run(p, "", size=10)
                continue

            checked = _matches(value, opt["label"])
            mark = CHECKED if checked else UNCHECKED
            mark_color = GREEN if checked else MID_GRAY

            _add_run(p, f"{mark}  ", bold=True, size=11, color=mark_color)
            _add_run(p, opt["label"], size=10,
                     color=GRAY if checked else MID_GRAY,
                     bold=checked)
            if opt.get("sub"):
                _add_run(p, f"  — {opt['sub']}", italic=True, size=8,
                         color=RGBColor(0x88, 0x88, 0x88))


def _hours_table(doc, fields):
    """Day-by-day hours table populated from extracted fields."""
    days = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
    table = doc.add_table(rows=len(days)+1, cols=5)
    table.style = "Table Grid"

    # Header
    headers = ["Day", "Open Time", "Close Time", "Closed", "By Appt"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        _set_cell_bg(cell, "F0F4FA")
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        _add_run(p, h, bold=True, size=9, color=BLUE)

    for r, day in enumerate(days):
        key = f"hours_{day.lower()}"
        h = fields.get(key)
        bg = "FFFFFF" if r % 2 == 0 else "F7F9FC"

        # Day name
        cell = table.cell(r+1, 0)
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], day, size=9, color=GRAY)

        # Open / Close
        o_val = ""
        c_val = ""
        closed = False
        by_appt = False

        if isinstance(h, dict):
            closed  = h.get("closed", False)
            by_appt = h.get("by_appt", False)
            o_val   = str(h.get("open", "")) if not closed and not by_appt else ""
            c_val   = str(h.get("close", "")) if not closed and not by_appt else ""
        elif isinstance(h, str) and h:
            o_val = h

        for col_idx, val in enumerate([o_val, c_val]):
            cell = table.cell(r+1, col_idx+1)
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            _add_run(cell.paragraphs[0], val, size=9, color=GRAY)

        # Closed checkbox
        cell = table.cell(r+1, 3)
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        mark = CHECKED if closed else UNCHECKED
        _add_run(cell.paragraphs[0], f"{mark}  Closed",
                 size=9, color=GREEN if closed else MID_GRAY, bold=closed)

        # By Appt checkbox
        cell = table.cell(r+1, 4)
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        mark = CHECKED if by_appt else UNCHECKED
        _add_run(cell.paragraphs[0], f"{mark}  By Appt",
                 size=9, color=GREEN if by_appt else MID_GRAY, bold=by_appt)


def _services_table(doc, services):
    """Populated services/products table."""
    if not services or not isinstance(services, list):
        _write_in(doc, None, "No services provided")
        return

    table = doc.add_table(rows=len(services)+1, cols=3)
    table.style = "Table Grid"

    headers = ["Service / Product", "Description", "Price"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        _set_cell_bg(cell, "F0F4FA")
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], h, bold=True, size=9, color=BLUE)

    for r, svc in enumerate(services):
        bg = "FFFFFF" if r % 2 == 0 else "F7F9FC"
        if isinstance(svc, dict):
            vals = [_safe(svc.get("name")), _safe(svc.get("description")), _safe(svc.get("price"))]
        else:
            vals = [str(svc), "—", "—"]
        for c, val in enumerate(vals):
            cell = table.cell(r+1, c)
            _set_cell_bg(cell, bg)
            cell.paragraphs[0].clear()
            _add_run(cell.paragraphs[0], val, size=9, color=GRAY)


def _faq_table(doc, fields):
    """Populated FAQ table."""
    faqs = [
        ("Do you take walk-ins?", fields.get("faq_walk_ins")),
        ("What forms of payment?", fields.get("faq_payment_methods")),
        ("Cancellation policy?", fields.get("faq_cancellation_policy")),
        ("Taking new customers?", fields.get("faq_new_customers")),
        ("Parking / location?", fields.get("faq_parking_location")),
    ]
    extra = fields.get("faq_additional", [])
    if isinstance(extra, list):
        for item in extra:
            if isinstance(item, dict):
                faqs.append((item.get("q", "Additional"), item.get("a", "")))
            elif isinstance(item, str) and item:
                faqs.append(("Additional", item))

    filled = [(q, a) for q, a in faqs if a and _safe(a) != "—"]
    if not filled:
        _write_in(doc, None, "No FAQ answers provided")
        return

    table = doc.add_table(rows=len(filled)+1, cols=2)
    table.style = "Table Grid"
    for i, h in enumerate(["Question", "Answer"]):
        cell = table.cell(0, i)
        _set_cell_bg(cell, "F0F4FA")
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], h, bold=True, size=9, color=BLUE)

    for r, (q, a) in enumerate(filled):
        bg = "FFFFFF" if r % 2 == 0 else "F7F9FC"
        cell = table.cell(r+1, 0)
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], str(q), size=9, color=BLUE, bold=True)

        cell = table.cell(r+1, 1)
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], _safe(a), size=9, color=GRAY)


# ─── MAIN BUILDER ────────────────────────────────────────────────────────────

def generate_intake_form(fields: dict, metadata: dict) -> bytes:
    """
    Build and return populated ARIA Client Intake Form as .docx bytes.
    Mirrors the structure of the original intake form exactly.
    Checkboxes checked in green (■) for client selections,
    write-in fields filled with client's actual words.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    f = fields
    business    = f.get("business_legal_name") or f.get("dba_name") or "New Client"
    received    = metadata.get("received_at", datetime.utcnow().isoformat())[:10]
    call_id     = metadata.get("conversation_id", "—")

    # ── TITLE ────────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    _add_run(title, "ARIA CLIENT INTAKE FORM", bold=True, size=22, color=BLUE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(2)
    _add_run(sub, "AI Receptionist Setup & Configuration — Completed via Voice Interview",
             size=12, color=LIGHT_BLUE)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_after = Pt(12)
    _add_run(meta_p,
             f"Business: {business}  |  Call ID: {call_id}  |  Completed: {received}",
             italic=True, size=8, color=MID_GRAY)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — BUSINESS INFORMATION
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "1", "Business Information")

    _question_label(doc, "Business Legal Name", required=True)
    _write_in(doc, f.get("business_legal_name"))

    _question_label(doc, "Doing Business As (DBA)")
    _write_in(doc, f.get("dba_name"))

    # Two-col: owner + title
    t = doc.add_table(rows=1, cols=2)
    for col, (lbl, key) in enumerate([("Owner / Primary Contact Name", "owner_name"),
                                       ("Title / Role", "owner_title")]):
        cell = t.cell(0, col)
        cell.paragraphs[0].clear()
        _add_run(cell.paragraphs[0], lbl, bold=True, size=10, color=GRAY)
        cell.add_paragraph()
        p = cell.paragraphs[1]
        val = _safe(f.get(key))
        _add_run(p, val, size=10, color=GRAY if val != "—" else MID_GRAY,
                 italic=val == "—")

    _question_label(doc, "Phone Number ARIA Will Answer", required=True)
    _write_in(doc, f.get("owner_phone"))

    _question_label(doc, "Business Address", required=True)
    _write_in(doc, f.get("business_address"))

    _question_label(doc, "Business Website")
    _write_in(doc, f.get("website_url"))

    _question_label(doc, "Years in Operation")
    _write_in(doc, f.get("years_in_operation"))

    _question_label(doc, "Industry / Business Type", required=True)
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Barbershop / Hair Salon"},  {"label": "Nail Salon / Spa"},
        {"label": "Dental Office"},             {"label": "Medical / Urgent Care"},
        {"label": "Veterinary Clinic"},         {"label": "Chiropractic / Physical Therapy"},
        {"label": "Mental Health / Counseling"},{"label": "Auto Dealership"},
        {"label": "Auto Repair / Quick Lube"},  {"label": "Property Management"},
        {"label": "Real Estate Agency"},        {"label": "Restaurant / Food Service"},
        {"label": "Retail Store"},              {"label": "Fitness / Gym / Wellness"},
        {"label": "Legal / Law Office"},        {"label": "Accounting / Financial Services"},
        {"label": "Insurance Agency"},          {"label": "Cleaning / Janitorial"},
        {"label": "Landscaping / Lawn Care"},   {"label": "Pool Service"},
        {"label": "Plumbing / HVAC / Electrical"},{"label": "Roofing / Construction"},
        {"label": "Photography / Videography"}, {"label": "Event Planning / Catering"},
        {"label": "Pest Control"},              {"label": "Other — describe below"},
    ], f.get("business_type"))
    _write_in(doc, f.get("business_description"), "Business description / additional type details")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — CONTACT & DELIVERY
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "2", "Contact & Report Delivery")

    _question_label(doc, "Email Address for Call Summaries & Recordings", required=True)
    emails = f.get("report_email_addresses")
    _write_in(doc, emails)

    _question_label(doc, "Best Phone to Reach You", required=True)
    _write_in(doc, f.get("owner_phone"))

    _question_label(doc, "How would you like to receive post-call reports?")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Email — Full Summary + MP3 Recording", "sub": "default"},
        {"label": "Email — Summary Only"},
        {"label": "Daily Digest — one email with all calls"},
        {"label": "Immediate — after every single call"},
        {"label": "Only notify me for calls needing follow-up"},
        {"label": "Only notify me for missed or unanswered calls"},
        {"label": "SMS / Text notification with link", "sub": "future feature"},
        {"label": "No preference — default settings are fine"},
    ], f.get("report_delivery_preference"))
    _write_in(doc, None, "Additional delivery notes")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — HOURS
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "3", "Hours of Operation")
    _hours_table(doc, f)
    _question_label(doc, "Holiday Closures / Seasonal Changes")
    _write_in(doc, f.get("holiday_notes") or f.get("seasonal_variations"), lines=2)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — GREETING & PERSONALITY
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "4", "ARIA Greeting & Personality")

    _question_label(doc, "What name should ARIA use?", required=True)
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "ARIA", "sub": "InSync Tech default"},
        {"label": "Use my business name only"},
        {"label": "A custom agent name", "sub": "see below"},
        {"label": "No preference — you decide"},
    ], f.get("agent_name"))
    _write_in(doc, f.get("agent_name"), "Custom agent name")

    _question_label(doc, "Opening Greeting Style", required=True)
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "\"Thank you for calling [Business], this is ARIA...\""},
        {"label": "\"Hi! You've reached [Business]...\"", "sub": "casual & friendly"},
        {"label": "\"Good [morning/afternoon], [Business]...\"", "sub": "time-aware"},
        {"label": "\"Welcome to [Business], how can I help you?\""},
        {"label": "Custom — I'll write my own script below"},
        {"label": "No preference — you decide"},
    ], f.get("greeting_script_raw"))
    _question_label(doc, "Custom Greeting Script")
    _write_in(doc, f.get("greeting_script_raw"), lines=3)

    _question_label(doc, "Personality & Tone")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Warm & Welcoming", "sub": "friendly, approachable"},
        {"label": "Professional & Polished", "sub": "formal, corporate"},
        {"label": "Friendly & Conversational", "sub": "casual, upbeat"},
        {"label": "Brief & Direct", "sub": "efficient, no filler"},
        {"label": "Empathetic & Patient", "sub": "ideal for medical/wellness"},
        {"label": "Energetic & Enthusiastic", "sub": "ideal for fitness/retail"},
        {"label": "Calm & Reassuring", "sub": "ideal for legal/counseling"},
        {"label": "No preference — you decide"},
    ], f.get("tone_preference"))

    _question_label(doc, "Voice Preference")
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "Female"},
        {"label": "Male"},
        {"label": "Neutral / Androgynous"},
        {"label": "No preference"},
    ], f.get("voice_gender_preference"))

    _question_label(doc, "Language Support")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "English Only"},
        {"label": "Spanish"},
        {"label": "English + Spanish (bilingual)"},
        {"label": "Other — write below"},
    ], f.get("ai_disclosure_preference"))
    _write_in(doc, None, "Additional language preferences")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — CALL HANDLING
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "5", "Call Handling & Behavior")

    _question_label(doc, "Primary purpose of most incoming calls", required=True)
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Schedule / Book an Appointment"},
        {"label": "Ask About Hours or Location"},
        {"label": "Get Pricing Information"},
        {"label": "Ask About Services Offered"},
        {"label": "Place or Check an Order"},
        {"label": "Request a Callback or Estimate"},
        {"label": "Cancel or Reschedule an Appointment"},
        {"label": "Check Job / Service Status"},
        {"label": "Billing or Payment Question"},
        {"label": "Emergency or Urgent Issue"},
        {"label": "New Client / New Patient Inquiry"},
        {"label": "Referral or Partner Inquiry"},
        {"label": "Complaint or Concern"},
        {"label": "General Question"},
    ], f.get("primary_call_reasons"))
    _write_in(doc, f.get("common_repetitive_questions"), "Other call types / common questions")

    _question_label(doc, "Information ARIA should capture from EVERY caller", required=True)
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Full Name"},
        {"label": "Callback Phone Number"},
        {"label": "Email Address"},
        {"label": "Reason for Calling"},
        {"label": "Preferred Callback Time"},
        {"label": "Service or Product of Interest"},
        {"label": "Appointment Date / Time Requested"},
        {"label": "New or Existing Client?"},
        {"label": "How They Heard About You"},
        {"label": "Urgency Level"},
        {"label": "Insurance Provider", "sub": "medical/dental"},
        {"label": "Vehicle Year / Make / Model", "sub": "auto"},
        {"label": "Property Address", "sub": "real estate/property mgmt"},
        {"label": "Pet Name & Species", "sub": "veterinary"},
    ], f.get("required_caller_fields"))
    _write_in(doc, f.get("industry_specific_capture_fields"), "Other fields to capture")

    _question_label(doc, "How should ARIA handle appointment requests?")
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "Capture details — I'll call back to confirm"},
        {"label": "Direct to my online booking link", "sub": "see below"},
        {"label": "Read available time slots I provide"},
        {"label": "Direct to call back during business hours"},
        {"label": "Take a voicemail — I'll follow up"},
        {"label": "Walk-ins only — we don't take appointments"},
    ], f.get("appointment_handling_method"))
    _write_in(doc, f.get("booking_link"), "Online booking link")

    _question_label(doc, "After-Hours Handling", required=True)
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "State hours and offer to take a message"},
        {"label": "State hours and direct to online booking"},
        {"label": "Take name & number for next-day callback"},
        {"label": "Play a custom after-hours message", "sub": "see below"},
        {"label": "Forward to emergency number", "sub": "see below"},
        {"label": "Same handling as business hours"},
    ], f.get("after_hours_script"))
    _write_in(doc, f.get("after_hours_script"), "Custom after-hours message / emergency number", lines=2)

    _question_label(doc, "Escalation — if a caller urgently needs a human:", required=True)
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "Offer to take a message for callback"},
        {"label": "Transfer to a specific number", "sub": "see below"},
        {"label": "Provide a direct callback number", "sub": "see below"},
        {"label": "Direct to voicemail"},
        {"label": "Emergency contact only for true emergencies"},
        {"label": "Tell them to call back during business hours"},
    ], f.get("escalation_path"))
    _write_in(doc, f.get("escalation_number"), "Escalation number / instructions")

    _question_label(doc, "Calls ARIA should NEVER handle — always escalate immediately")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Legal threats or mentions of lawsuits"},
        {"label": "Extremely upset or hostile callers"},
        {"label": "Media or press inquiries"},
        {"label": "Insurance claims or disputes"},
        {"label": "Any mention of injury or medical emergency"},
        {"label": "Billing disputes above a certain amount"},
        {"label": "VIP or high-value clients"},
        {"label": "Vendor or supplier calls"},
        {"label": "Government or regulatory inquiries"},
        {"label": "Other — describe below"},
    ], f.get("escalation_trigger_list"))
    _write_in(doc, f.get("complaint_handling"), "Additional escalation notes")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — SERVICES & KNOWLEDGE
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "6", "Services, Products & Business Knowledge")

    _question_label(doc, "Main services / products ARIA should know about", required=True)
    _services_table(doc, f.get("services_list"))

    _question_label(doc, "How should ARIA handle pricing questions?")
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "Share exact prices I provide above"},
        {"label": "Share price ranges only"},
        {"label": "Pricing varies — direct to quote or consultation"},
        {"label": "Never discuss pricing — direct to me"},
        {"label": "Direct to website for pricing"},
        {"label": "Answer per-service based on what I provide"},
    ], f.get("pricing_approach"))
    _write_in(doc, f.get("pricing_details"), "Pricing details / exceptions")

    _question_label(doc, "Frequently Asked Questions")
    _faq_table(doc, f)

    _question_label(doc, "Things ARIA should NEVER say or promise")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Never quote exact prices without my approval"},
        {"label": "Never promise same-day availability"},
        {"label": "Never make medical, legal, or financial claims"},
        {"label": "Never discuss active complaints or lawsuits"},
        {"label": "Never mention competitor names"},
        {"label": "Never discuss staffing or employee matters"},
        {"label": "Never discuss internal business operations"},
        {"label": "Other — describe below"},
    ], f.get("prohibited_statements"))
    _write_in(doc, f.get("prohibited_statements"), "Additional prohibited statements")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — NOTIFICATIONS
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "7", "Notifications & Follow-Up")

    _question_label(doc, "Which calls should trigger a notification?")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Every call — no exceptions"},
        {"label": "Only calls where caller left a message"},
        {"label": "Only calls requesting a callback"},
        {"label": "Only calls requesting an appointment"},
        {"label": "Only calls mentioning a complaint"},
        {"label": "Only new caller inquiries"},
        {"label": "Only missed or unanswered calls"},
        {"label": "Only calls flagged as urgent"},
        {"label": "Weekly summary digest only"},
        {"label": "No notifications — I'll check manually"},
    ], f.get("notification_triggers"))

    _question_label(doc, "Should ARIA send automated follow-up to callers?")
    _check_banner(doc, include_check_all=False)
    _check_grid(doc, [
        {"label": "No — I'll follow up manually"},
        {"label": "Yes — follow-up text to caller", "sub": "if number captured"},
        {"label": "Yes — follow-up email to caller", "sub": "if email captured"},
        {"label": "Yes — appointment confirmation message"},
        {"label": "Not now but interested in this feature"},
    ], f.get("followup_preference"))
    _write_in(doc, f.get("future_followup_interest"), "Follow-up notes")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — CURRENT SOFTWARE & SYSTEMS
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "9", "Current Software & Systems")

    _question_label(doc, "Phone System / Current Call Handling")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Standard landline"},         {"label": "VoIP (RingCentral, Vonage, 8x8)"},
        {"label": "Mobile only"},               {"label": "Google Voice"},
        {"label": "Grasshopper"},               {"label": "OpenPhone"},
        {"label": "Nextiva"},                   {"label": "No formal system — just a cell"},
        {"label": "Other — describe below"},    {"label": ""},
    ], f.get("current_phone_system"))
    _write_in(doc, f.get("call_routing_setup"), "Phone system / call routing details")

    _question_label(doc, "Scheduling & Booking Software")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "No system — manual / phone only"},  {"label": "Calendly"},
        {"label": "Acuity Scheduling"},                {"label": "Google Calendar"},
        {"label": "Square Appointments"},              {"label": "Booksy", "sub": "salons / barbers"},
        {"label": "Jane App", "sub": "healthcare"},    {"label": "Mindbody", "sub": "fitness"},
        {"label": "SimplePractice", "sub": "mental health"}, {"label": "Vagaro", "sub": "beauty"},
        {"label": "OpenDental / Dentrix", "sub": "dental"},  {"label": "Epic / Athenahealth", "sub": "medical"},
        {"label": "ResMan / AppFolio", "sub": "property mgmt"}, {"label": "CDK / Reynolds", "sub": "auto"},
        {"label": "Other — describe below"},           {"label": ""},
    ], f.get("scheduling_software"))
    _write_in(doc, f.get("booking_workflow"), "Scheduling system / booking workflow")

    _question_label(doc, "CRM & Lead Management")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "No CRM — manage manually"},  {"label": "HubSpot"},
        {"label": "Salesforce"},                {"label": "Zoho CRM"},
        {"label": "Pipedrive"},                 {"label": "GoHighLevel"},
        {"label": "Keap / Infusionsoft"},       {"label": "Monday.com"},
        {"label": "Airtable"},                  {"label": "Google Sheets", "sub": "informal CRM"},
        {"label": "Follow Up Boss", "sub": "real estate"}, {"label": "Other — describe below"},
    ], f.get("crm_software"))
    _write_in(doc, f.get("customer_data_location"), "CRM / where customer data lives today")

    _question_label(doc, "Point of Sale & Payment Processing")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "No POS — cash or manual"},  {"label": "Square"},
        {"label": "Stripe"},                   {"label": "Clover"},
        {"label": "Toast", "sub": "restaurants"}, {"label": "Lightspeed"},
        {"label": "Shopify POS"},              {"label": "PayPal / Venmo"},
        {"label": "QuickBooks Payments"},      {"label": "Other — describe below"},
        {"label": ""},                         {"label": ""},
    ], f.get("payment_system"))
    _write_in(doc, f.get("online_payment_capability"), "Payment system details")

    _question_label(doc, "Email, Marketing & Reviews")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "No email marketing"},        {"label": "Mailchimp"},
        {"label": "Constant Contact"},          {"label": "Klaviyo"},
        {"label": "ActiveCampaign"},            {"label": "GoHighLevel"},
        {"label": "Google Business Profile"},   {"label": "Meta Business Suite"},
        {"label": "Yelp for Business"},         {"label": "Birdeye / Podium / Broadly"},
        {"label": "Other — describe below"},    {"label": ""},
    ], f.get("email_marketing_tools"))
    _write_in(doc, f.get("review_platforms"), "Marketing / review platforms used")

    _question_label(doc, "Accounting & Operations")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "No accounting software"},   {"label": "QuickBooks Online"},
        {"label": "QuickBooks Desktop"},       {"label": "FreshBooks"},
        {"label": "Xero"},                     {"label": "Wave"},
        {"label": "Gusto", "sub": "payroll"},  {"label": "ADP", "sub": "payroll"},
        {"label": "Jobber", "sub": "field service"}, {"label": "ServiceTitan", "sub": "trades"},
        {"label": "Other — describe below"},   {"label": ""},
    ], f.get("accounting_software"))
    _write_in(doc, f.get("operations_software"), "Operations / accounting details")

    _question_label(doc, "Integration Wishlist — systems you'd want ARIA to connect to")
    _write_in(doc, f.get("integration_wishlist"), "e.g. Push callers into HubSpot / Check Square calendar for availability", 3)

    _question_label(doc, "Highest-Value Integration")
    _write_in(doc, f.get("highest_value_integration"), "The single connection that would change your day most")

    _question_label(doc, "Workflow Pain Points")
    _write_in(doc, f.get("workflow_pain_points"), "Current tech frustrations / inefficiencies", 3)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — FUTURE FEATURES
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "10", "Future Features & Wishlist")

    _question_label(doc, "What would you like your ARIA agent to do in the future?")
    _check_banner(doc, include_check_all=True)
    _check_grid(doc, [
        {"label": "Two-way SMS texting with callers"},
        {"label": "Outbound reminder calls or texts to clients"},
        {"label": "Appointment reminders — sent automatically"},
        {"label": "Calendar / booking integration", "sub": "Calendly, Google, etc."},
        {"label": "Live call transfer to a staff member"},
        {"label": "CRM integration", "sub": "HubSpot, Salesforce, etc."},
        {"label": "Email auto-response to callers"},
        {"label": "Payment collection over the phone"},
        {"label": "Intake forms sent to caller after the call"},
        {"label": "Multi-location / multi-line support"},
        {"label": "Full bilingual AI support"},
        {"label": "Analytics dashboard — call trends & volume"},
        {"label": "Custom branded caller ID / number"},
        {"label": "After-hours emergency escalation chain"},
        {"label": "Google Reviews request sent after call"},
        {"label": "Chatbot version for my website"},
        {"label": "ARIA handling social media DMs or emails"},
        {"label": "Integration with my existing phone system"},
        {"label": "Industry-specific software integration", "sub": "dental, property mgmt, etc."},
        {"label": "White-label ARIA for my own clients", "sub": "agencies / resellers"},
        {"label": "ARIA acting as a full virtual office manager"},
        {"label": "Something else — describe below"},
    ], f.get("wishlist_features"))
    _write_in(doc, f.get("wishlist_features") if isinstance(f.get("wishlist_features"), str) else None,
              "Additional feature requests")
    _write_in(doc, f.get("pain_points_remaining"), "Remaining pain points / suggestions")
    _write_in(doc, f.get("future_vision"), "Future vision notes")

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — AUTHORIZATION (pre-filled from interview)
    # ══════════════════════════════════════════════════════════════════════════
    _section_header(doc, "9", "Authorization")

    auth_p = doc.add_paragraph()
    auth_p.paragraph_format.space_before = Pt(6)
    auth_p.paragraph_format.space_after = Pt(10)
    _add_run(auth_p,
             "By completing this voice interview, the client has confirmed that the information "
             "provided is accurate and authorized the configuration of AI receptionist services "
             "for their business by InSync Tech.", size=10, color=GRAY)

    # Pre-filled authorization fields
    for label_text, val in [
        ("Authorized By (Name):", f.get("owner_name")),
        ("Business:", business),
        ("Interview Date:", received),
        ("Call ID:", call_id),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _add_run(p, f"{label_text}  ", bold=True, size=10, color=MID_GRAY)
        _add_run(p, _safe(val), size=10, color=GRAY)

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(16)
    _add_run(footer_p,
             "InSync Tech Inc.  |  Venice, Florida  |  insynctech.io  |  "
             "Completed via ARIA Voice Interview — Internal Use Only",
             italic=True, size=8, color=MID_GRAY)

    # ── SAVE TO BYTES ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
